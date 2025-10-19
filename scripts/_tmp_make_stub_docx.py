from docx import Document
p = Document()
p.add_heading("Free Datasets (Stub)", level=1)
p.add_paragraph("NIMH Anxiety: https://www.nimh.nih.gov/health/topics/anxiety-disorders")
p.add_paragraph("MedlinePlus Depression: https://medlineplus.gov/depression.html")
p.add_paragraph("WHO mhGAP: https://www.who.int/teams/mental-health-and-substance-use/policy-law-rights/mhgap")
p.save("data/raw/lists/List - Free Datasets - SukoonAI.docx")
print("Wrote stub DOCX → data/raw/lists/List - Free Datasets - SukoonAI.docx")
